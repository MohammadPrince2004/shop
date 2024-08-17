#############################################



##################################################
""" ------------ a system for any shop --------"""
##################################################


######################################
"""---------- The modules ----------"""


from customtkinter import *
from PIL import Image
import requests
import json
from threading import Thread
import cv2
from pyzbar.pyzbar import decode
import time
import os
import numpy as np
import tkinter
import imutils
import qrcode
from PIL import ImageDraw
from PIL import ImageFont
import docx
from datetime import date
################################################

##################################################

set_appearance_mode("dark")
set_default_color_theme("green")
img=CTkImage(dark_image=Image.open(r"F:\Downloads\Downloaded\Downloads\Shop App.png"),size=(150,100))
img2=CTkImage(dark_image=Image.open(r"F:\Downloads\Downloaded\Downloads\bell-line-icon (1).png"),size=(30,30))
img3=CTkImage(dark_image=Image.open(r"F:\bud.png"),size=(50,50))
img4=CTkImage(dark_image=Image.open(r"C:\Users\moh\new.png"),size=(200,200))
#######################################
"""----------------- the main class ------------"""

class shop(CTk):

    def __init__(self):
        super().__init__()
        self.title("shop")
        self.geometry("800x600+500+200")
        self.iconbitmap(r"F:\Downloads\Downloaded\Downloads\onlineshop_78377.ico")
        self.q="go"
        self.name=""
        self.tot_price=0
        self.b_price=0
        self.time=time.asctime()
        self.v_code=StringVar()
        self.v_price=StringVar()
        self.v_name=StringVar()
        self.new_person={}
        self.ls_bts=[]
        self.ls_earns=[]
        self.ls_fatora={}
        self.ls_widgets=[[],[],[],[],[],[],[],[],[],[],[]]
        self.grid_rowconfigure(0,weight=1)
        self.grid_columnconfigure(0,weight=0)
        self.grid_columnconfigure(1,weight=1)
        self.indx=-1

        #frame of the notifications and scrollable frame of any choice

        self.frame_right=CTkFrame(self,corner_radius=0)
        self.frame_right.grid(row=0,column=1,sticky="nswe",padx=5,pady=10)
        self.frame_right.grid_rowconfigure(0,weight=0)
        self.frame_right.grid_rowconfigure(1,weight=1)
        self.frame_right.grid_columnconfigure(0,weight=1)
        
        self.frame_left=CTkFrame(self,corner_radius=0)
        self.frame_left.grid(row=0,column=1)
        self.frame_left.grid(row=0,column=0,sticky="nswe",padx=5,pady=10)
        self.frame_left.grid_rowconfigure(0,weight=1)
        self.frame_left.grid_rowconfigure((1,2,3,4,5,6,7,8,9),weight=0)
        self.frame_left.grid_rowconfigure(10,weight=1)
        self.frame_logo_shop=CTkFrame(self.frame_left,fg_color="transparent")
        self.frame_logo_shop.grid(row=0,column=0)        

        self.logo_label=CTkLabel(self.frame_logo_shop,text="",image=img)
        self.logo_label.grid(row=0,column=0,sticky="nswe")

        self.bt1=CTkButton(self.frame_left,text=" add any goad",fg_color="#353535",command=self.add_com,hover_color="#1f685a",corner_radius=0)
        self.bt1.grid(row=1,column=0,padx=5,pady=10)
        self.bt2=CTkButton(self.frame_left,text=" Modifing the prices",fg_color="#353535",command=self.modfy_com,hover_color="#1f685a",corner_radius=0)
        self.bt2.grid(row=2,column=0,padx=5,pady=0)
        self.bt3=CTkButton(self.frame_left,text=" About prices",fg_color="#353535",command=self.prices_com,hover_color="#1f685a",corner_radius=0)
        self.bt3.grid(row=3,column=0,padx=5,pady=10)
        self.bt4=CTkButton(self.frame_left,text=" Limits of Prices",fg_color="#353535",command=self.limits_com,hover_color="#1f685a",corner_radius=0)
        self.bt4.grid(row=4,column=0,padx=5,pady=0)
        self.bt5=CTkButton(self.frame_left,text=" Creating a fatora",fg_color="#353535",command=self.creating_fatora_com,hover_color="#1f685a",corner_radius=0)
        self.bt5.grid(row=5,column=0,padx=5,pady=10)        
        self.bt6=CTkButton(self.frame_left,text=" Back any goad",fg_color="#353535",command=self.back_goad,hover_color="#1f685a",corner_radius=0)
        self.bt6.grid(row=6,column=0,padx=5,pady=0)
        self.bt7=CTkButton(self.frame_left,text=" Creating a user "+"\n"+" new account ",fg_color="#353535",command=self.newuser_com,hover_color="#1f685a",corner_radius=0)
        self.bt7.grid(row=7,column=0,padx=5,pady=10)
        self.bt8=CTkButton(self.frame_left,text=" Searching any fatora ",fg_color="#353535",command=self.search_fatora_com,hover_color="#1f685a",corner_radius=0)
        self.bt8.grid(row=8,column=0,padx=5,pady=0)
        self.bt9=CTkButton(self.frame_left,text=" Searching about "+"\n"+"persons",fg_color="#353535",command=self.srch_persons_com,hover_color="#1f685a",corner_radius=0)
        self.bt9.grid(row=9,column=0,padx=5,pady=10)

        self.bt_qr=CTkButton(self.frame_left,text="Generate a QR-code",fg_color="#353535",command=self.gen_qr_code,hover_color="#1f685a",corner_radius=0)
        self.bt_qr.grid(row=10,column=0,pady=0,sticky="n")

        self.bt10=CTkButton(self.frame_left,text="About Us",fg_color="#353535",command=self.aboutus_com,hover_color="#1f685a",corner_radius=0)
        self.bt10.grid(row=10,column=0,padx=5,pady=10,sticky="s")
        
        self.upper_frame=CTkFrame(self.frame_right,fg_color="#1f685a",corner_radius=0)
        self.upper_frame.grid(row=0,column=0,sticky="nswe")
        
        self.upper_frame.grid_columnconfigure((1,2),weight=0)
        self.upper_frame.grid_columnconfigure(0,weight=1)
        self.upper_frame.grid_rowconfigure(0,weight=1)
        self.create=False
        self.back=False

        self.not_budget_frame=CTkFrame(self.frame_right)
        
        for i in range(11):
            self.sc_frame=CTkScrollableFrame(self.frame_right,corner_radius=5)
            # self.sc_frame.grid(row=1,column=0,sticky="nswe")
            self.ls_bts.append(self.sc_frame)

        self.noti_frame=CTkFrame(self.upper_frame,fg_color="transparent")
        self.noti_frame.grid(row=0,column=2,padx=10 ,pady=2)
        self.noti_frame.grid_rowconfigure(0,weight=0)
        self.noti_frame.grid_rowconfigure(1,weight=1)
        self.noti_frame.grid_columnconfigure((0,2),weight=0)
        self.noti_frame.grid_columnconfigure(1,weight=1)

        self.logo_label_notiup=CTkLabel(self.noti_frame,text="   ",corner_radius=5,width=5,height=5)# the logo of red
        self.logo_label_notiup.grid(row=0,column=2,padx=2)
        self.logo_not=CTkLabel(self.noti_frame,text="",image=img2)
        self.logo_not.grid(row=1,column=1)
        # self.name_not=CTkLabel(self.noti_frame,text="Notifications",font=("arial",14,"bold"))
        # self.name_not.grid(row=1,column=0,padx=5,)
        self.logo_not.bind("<Button-1>",self.not_com)


        self.budget_frame=CTkFrame(self.upper_frame,fg_color="transparent")
        self.budget_frame.grid(row=0,column=0,padx=5,pady=5,sticky="w")
        self.budget_frame.grid_rowconfigure(0,weight=1)
        
        self.budget_frame.grid_columnconfigure(0,weight=1)
        self.budget_frame.grid_columnconfigure(1,weight=0)
        # self.l=CTkLabel(self.budget_frame,text="",height=5)
        # self.l.grid(row=0,column=1)

        self.label_logo_budget=CTkLabel(self.budget_frame,image=img3,text="")
        self.label_logo_budget.grid(row=0,column=0,pady=5,padx=20,sticky="w")
        self.name_budget=CTkLabel(self.budget_frame,text="0",font=("arial",14,"bold"))
        self.name_budget.grid(row=0,column=1,padx=5,pady=1)
        self.open_close_not=False
        self.b=False
        self.thread()
        self.thread_show()
       
     
        
    def create_earns_json(self):
        try:
            data={}

            if os.path.exists(r"F:\earns.json"):

                with open(r"F:\earns.json","r") as f:

                    data=json.load(f)
                if str(date.today()) in data:
                    pass
                else:
                    data[str(date.today())]=0
            else:
                data[str(date.today())]=0

                with open(r"F:\earns.json","w") as f:

                    json.dump(data,f)

            if self.back:
                    # to back any goad

                data[str(date.today())]=int(data[str(date.today())])-self.b_price

            elif self.create:

                    # when creating a fatora
                data[str(date.today())]=data[str(date.today())]+self.tot_price

            with open(r"F:\earns.json","w") as f:

                json.dump(data,f)
        except:
            pass


    def show_earns(self):

        data={}
        while True:

            self.create_earns_json()

            with open(r"F:\earns.json","r") as f:
                data=json.load(f)
            
            self.name_budget.configure(text=f"{data[str(date.today())]}")



        
    def thread_show(self):
        try:
            t1=Thread(target=self.show_earns)
            t1.start()
        except:
            pass

    
    def not_com(self,event):
        try:
            if not(self.b):

                self.b=True
                self.not_budget_frame.grid(row=0,column=1,rowspan=3,sticky="nswe")
                self.not_budget_frame.grid_columnconfigure(0,weight=1)
                self.not_budget_frame.grid_rowconfigure(0,weight=0)
                self.not_budget_frame.grid_rowconfigure(1,weight=1)
                self.bud_frame=CTkFrame(self.not_budget_frame,fg_color="blue",corner_radius=0)
                self.bud_frame.grid(row=0,column=0,sticky="we")
                self.notif_frame=CTkScrollableFrame(self.not_budget_frame)
                self.notif_frame.grid(row=1,column=0,sticky="nswe")
                self.notif_frame.grid_columnconfigure(0,weight=1)
                self.bud_frame.grid_rowconfigure(0,weight=1)
                self.bud_frame.grid_columnconfigure(0,weight=1)
                self.bud_frame.grid_columnconfigure(1,weight=0)
                self.not_label=CTkLabel(self.bud_frame,text="The goads of less amount",font=("arial",14,"bold"))
                self.not_label.grid(row=0,column=0,sticky="w",padx=5,pady=1)

                self.exit_bt=CTkButton(self.bud_frame,text="X",width=15,height=15,command=self.exit,fg_color="red")
                self.exit_bt.grid(row=0,column=1,sticky="e",pady=1)

                t1= Thread(target=self.show_not)
                t1.start()
        except:
            pass
    def exit(self):
        try:
            self.not_budget_frame.grid_forget()
            self.b=False
        except:
            pass
  
    def create_json_not(self):
        try:
            data={}
            data_not={}
            if os.path.exists(r"F:\notifications.json"):
                with open(r"F:\notifications.json","r") as f:
                    data_not=json.load(f)

            if os.path.exists(r"F:\goads.json"):
                with open(r"F:\goads.json","r") as f:
                    data=json.load(f)
                ls_info=list(data.values())
                ls_codes=list(data.keys())

                for i in range(len(ls_info)):
                    if len(ls_info[i])==4:
                        if int(ls_info[i][-1])==int(ls_info[i][0]):
                            data_not[ls_codes[i]]=ls_info[i][2]
                        elif int(ls_info[i][-1])<int(ls_info[i][0]):
                            if ls_codes[i] in data_not:
                                data_not.pop(ls_codes[i])
                            else:
                                pass

            with open(r"F:\notifications.json","w") as f:
                json.dump(data_not,f)
        except:
            pass

    def label_color(self):
        try:
            while True:
                data={}
                self.create_json_not()

                with open (r"F:\notifications.json","r") as f:
                    data=json.load(f)
                if len(data)==0:
                    self.logo_label_notiup.configure(fg_color="transparent")

                else:
                    self.logo_label_notiup.configure(fg_color="red")
        except:
            pass
            

    def thread(self):
        try:
            t1=Thread(target=self.label_color)
            t1.start()
        except:
            pass
   
    def show_not(self):
        try:
            self.create_json_not()
            data={}
            c=0
            with open(r"F:\notifications.json","r") as f:
                data=json.load(f)
            ls_info=list(data.values())
            for i in range(len(ls_info)):

                self.frame=CTkFrame(self.notif_frame,height=30, corner_radius=0)
                self.frame.grid(row=c,column=0,sticky="nswe",pady=1)
                self.frame.grid_rowconfigure(0,weight=1)



                self.label=CTkLabel(self.frame,text=f"{ls_info[i]}")
                self.label.grid(row=0,column=0,sticky="we",pady=2,padx=5)



                self.frame=CTkFrame(self.notif_frame,height=3, fg_color="blue")
                self.frame.grid(row=c+1,column=0,sticky="we",pady=1)

                c=c+2
        except:
            pass

    def qrcodesccaner(self):
        try:
            url = "http://192.168.1.3:8080/shot.jpg"
            detector = cv2.QRCodeDetector()
            # While loop to continuously fetching data from the Url
            ls=self.q.split("_")

            while not((self.q.split("_")[0]).isnumeric() and (self.q.split("_"))[1].isalpha) :

                img_resp = requests.get(url)
                img_arr = np.array(bytearray(img_resp.content), dtype=np.uint8)
                img = cv2.imdecode(img_arr, -1)
                img = imutils.resize(img, width=500, height=500)

                # print(type(img))
                decocdeQR = decode(img)

                # print(decocdeQR)


                if len(decocdeQR)==0:
                    continue
                else:

                    self.q=(decocdeQR[0].data.decode('ascii'))
                    print(self.q)
                    # return(decocdeQR[0].data.decode('ascii'))
            self.v_code.set(self.q.split("_")[0])
            self.v_name.set(self.q.split("_")[1])
            data={}
            if os.path.exists(r"F:\goad.json"):
                with open(r"F:\goads.json", "r") as f:
                    data = json.load(f)
                self.v_price.set(str(data[self.v_code]))
        except:
            pass

            
        
            


    def add_com(self):
        try:
            self.q="go"
            data={}
            if os.path.exists(r"F:\goads.json") :
                pass
            else:
                with open(r"F:\goads.json","w") as f:
                    json.dump(data,f)
            if self.indx!=0:
                for i in self.ls_widgets[self.indx]:
                    i.grid_forget()
                self.indx=0
                for i in range(11):
                        self.ls_bts[i].grid_forget()
                self.ls_bts[self.indx].grid(row=1,column=0,sticky="nswe")
                self.ls_bts[self.indx].grid_rowconfigure((0,1,2,3,4),weight=0)
                self.ls_bts[self.indx].grid_columnconfigure(0,weight=1)
                self.label_Add=CTkLabel(self.ls_bts[self.indx],text="Add a new goads ")
                self.label_Add.grid(row=0,column=0)
                self.ls_widgets[self.indx].append(self.label_Add)
                self.entry_code=CTkEntry(self.ls_bts[self.indx],placeholder_text="Enter the code ",width=400,height=25,textvariable=self.v_code)
                self.entry_code.grid(row=1,column=0,pady=10)
                self.ls_widgets[self.indx].append(self.entry_code)
                self.entry_name_goad=CTkEntry(self.ls_bts[self.indx],placeholder_text="Enter the name of the goad ",width=400,height=25,textvariable=self.v_name)
                self.entry_name_goad.grid(row=2,column=0,pady=10)
                self.ls_widgets[self.indx].append(self.entry_name_goad)

                self.entry_price=CTkEntry(self.ls_bts[self.indx],placeholder_text="Enter the price",width=400,height=25,textvariable=self.v_price)
                self.entry_price.grid(row=3,column=0,pady=10)
                self.ls_widgets[self.indx].append(self.entry_price)

                self.entry_amount=CTkEntry(self.ls_bts[self.indx],placeholder_text="Enter the amount",width=400,height=25)
                self.entry_amount.grid(row=4,column=0,pady=10)
                self.ls_widgets[self.indx].append(self.entry_amount)

                self.bt_add_goad=CTkButton(self.ls_bts[self.indx],text="Add",command=self.add_bt)
                self.bt_add_goad.grid(row=5,column=0,pady=0)
                self.ls_widgets[self.indx].append(self.bt_add_goad)

                self.label_d=CTkLabel(self.ls_bts[self.indx],text="",font=("arial",15,"bold"))
                self.label_d.grid(row=6,column=0)
                self.ls_widgets[self.indx].append(self.label_d)

                t1=Thread(target=self.qrcodesccaner)
                t1.start()
        except:
            pass

            # print(self.q)
            # print(self.entry_code.get())
            
                
    def add_bt(self):
        try:
            # print(self.v_price.get())
            ls=self.q.split("_")

            price=self.entry_price.get()
            amount=self.entry_amount.get()
            self.entry_amount.delete(0,END)
            self.entry_price.delete(0,END)
            self.entry_name_goad.delete(0,END)
            self.entry_code.delete(0,END)

            if len(ls)==2:

                if (price.isdecimal() and amount.isdecimal()):
                    data={ls[0]:[int(amount),int(price),ls[1]]}
                    if os.path.exists(r"F:\goads.json"):
                        print("error")
                        with open (r"F:\goads.json","r") as file:
                            data=json.load(file)
                        ls_codes=list(data.keys())

                        if ls[0] in ls_codes:

                            data[ls[0]][0]=int(data[ls[0]][0])+int(amount)
                            data[ls[0]][1]=int(price)
                        else:
                            data[ls[0]]=[int(amount),int(price),ls[1]]
                            print(data)
                    with open(r"F:\goads.json","w") as f:
                        json.dump(data,f)
                    self.create_json_not()

                    self.label_d.configure(text="Done")
                else:
                    self.label_d.configure(text="Enter the data correctly !!")
            else:
                self.label_d.configure(text="Scan the QR-Codes to fill the upper entries !!")

            self.q="go"
            t1=Thread(target=self.qrcodesccaner)
            t1.start()
        except:
            pass

            
        


       
    def modfy_com(self):
        try:
            self.q="go"
            if self.indx!=1:
                for i in self.ls_widgets[self.indx]:
                    i.grid_forget()
                self.indx=1
                for i in range(11):
                        self.ls_bts[i].grid_forget()

                self.ls_bts[self.indx].grid(row=1,column=0,sticky="nswe")
                self.ls_bts[self.indx].grid_rowconfigure((0,1,2,3,4),weight=0)
                self.ls_bts[self.indx].grid_columnconfigure(0,weight=1)
                self.label_mod=CTkLabel(self.ls_bts[self.indx],text="Modify the prices of goads ")
                self.label_mod.grid(row=0,column=0)
                self.ls_widgets[self.indx].append(self.label_mod)
                self.entry_code=CTkEntry(self.ls_bts[self.indx],placeholder_text="Enter the code ",width=400,height=25,textvariable=self.v_code)
                self.entry_code.grid(row=1,column=0,pady=10)
                self.ls_widgets[self.indx].append(self.entry_code)
                self.entry_price=CTkEntry(self.ls_bts[self.indx],placeholder_text="Enter the price",width=400,height=25)
                self.entry_price.grid(row=2,column=0,pady=10)
                self.ls_widgets[self.indx].append(self.entry_price)
                self.bt_mod_goad=CTkButton(self.ls_bts[self.indx],text="Modify",command=self.mod_bt)
                self.bt_mod_goad.grid(row=3,column=0,pady=0)
                self.ls_widgets[self.indx].append(self.bt_mod_goad)
                self.label_d=CTkLabel(self.ls_bts[self.indx],text="",font=("arial",15,"bold"))
                self.label_d.grid(row=4,column=0)
                self.ls_widgets[self.indx].append(self.label_d)

                t1=Thread(target=self.qrcodesccaner)
                t1.start()
        except:
            pass

    def mod_bt(self):
        try:
            data={}
            ls=self.q.split("_")
            price=self.entry_price.get()
            self.entry_price.delete(0,END)
            # data={ls[0]:[int(amount),int(price),ls[1]]}
            if (price.isdecimal() and ls[0].isnumeric()):

                if os.path.exists(r"F:\goads.json"):
                    file=open(r"F:\goads.json")
                    data=json.load(file)
                    ls_codes=list(data.keys())
                    file.close()
                    if ls[0] in ls_codes:
                        data[ls[0]][1]=int(price)
                        self.label_d.configure(text="Done")
                        with open(r"F:\goads.json","w") as f:
                            json.dump(data,f)
                    else:
                        self.label_d.configure(text="Please add the goad fisrt !!")

                else:
                    self.label_d.configure(text="Please add the goad fisrt !!")
            else:
                self.label_d.configure(text="Enter the data correctly !!")


            self.q="go"
            t1=Thread(target=self.qrcodesccaner)
            t1.start()
        except:
            pass
        
    def prices_com(self):

        
        try:
            self.q="go"
            if self.indx!=2:
                for i in self.ls_widgets[self.indx]:
                    i.grid_forget()
                self.indx=2
                for i in range(11):
                        self.ls_bts[i].grid_forget()
                self.ls_bts[self.indx].grid(row=1,column=0,sticky="nswe")
                self.ls_bts[self.indx].grid_rowconfigure(0,weight=0)
                self.ls_bts[self.indx].grid_rowconfigure(1,weight=0)
                self.ls_bts[self.indx].grid_rowconfigure(2,weight=0)
                self.ls_bts[self.indx].grid_columnconfigure(0,weight=1)

                self.label_tit=CTkLabel(self.ls_bts[self.indx],text="Scan the QR-Code of the goad to get the price ",font=("arial",15,"bold"))
                self.label_tit.grid(row=0,column=0,sticky="we",pady=5)
                self.ls_widgets[self.indx].append(self.label_tit)

                self.entry_code=CTkEntry(self.ls_bts[self.indx],textvariable=self.v_code,placeholder_text="Enter the code")
                self.entry_code.grid(row=1,column=0,sticky="we",pady=5)
                self.ls_widgets[self.indx].append(self.entry_code)

                self.bt_about=CTkButton(self.ls_bts[self.indx],text="About",command=self.ab)
                self.bt_about.grid(row=2,column=0,pady=10)
                self.ls_widgets[self.indx].append(self.bt_about)

                self.label_ab=CTkLabel(self.ls_bts[self.indx],text="",font=("arial",15,"bold"))
                self.label_ab.grid(row=3,column=0,sticky="nswe")
                self.ls_widgets[self.indx].append(self.label_ab)

            
            
                t1=Thread(target=self.qrcodesccaner)
                t1.start()
        except:
            pass
            

           
    def ab(self):
        try:
            ls=self.q.split("_")
            self.q="go"
            file=open(r"F:\goads.json")
            data=json.load(file)
            file.close()
            if len(data)!=0:
                ls_codes=list(data.keys())
                ls_aboutgds=list(data.values())

                for i in ls_codes:

                    if i==ls[0]:

                        self.label_ab.configure(text="the price of the goad "+ls[1]+ " is "+str(data[i][1]))
                self.entry_code.delete(0,END)
            else:
                self.label_ab.configure(text="Please , Add the goad first !!")

            t1=Thread(target=self.qrcodesccaner)
            t1.start()
        except:
            pass

                
    def limits_com(self):
        # self.indx=3
        # # print(self.indx)
        # pass
       try:
            self.q="go"
            if self.indx!=3:
                for i in self.ls_widgets[self.indx]:
                    i.grid_forget()
                self.indx=3
                for i in range(11):
                    self.ls_bts[i].grid_forget()

                self.ls_bts[self.indx].grid(row=1,column=0,sticky="nswe")
                self.ls_bts[self.indx].grid_rowconfigure(0,weight=0)
                self.ls_bts[self.indx].grid_rowconfigure(1,weight=0)
                self.ls_bts[self.indx].grid_rowconfigure(2,weight=0)
                self.ls_bts[self.indx].grid_columnconfigure(0,weight=1)

                self.label_tit=CTkLabel(self.ls_bts[self.indx],text="Scan the QR-Code of the goad to Limit the num of goods ",font=("arial",15,"bold"))
                self.label_tit.grid(row=0,column=0,sticky="we",pady=5)
                self.ls_widgets[self.indx].append(self.label_tit)

                self.entry_code=CTkEntry(self.ls_bts[self.indx],textvariable=self.v_code,placeholder_text="Enter the code")
                self.entry_code.grid(row=1,column=0,sticky="we",pady=5)
                self.ls_widgets[self.indx].append(self.entry_code)

                self.entry_limit=CTkEntry(self.ls_bts[self.indx],placeholder_text="Enter the limit")
                self.entry_limit.grid(row=2,column=0,sticky="we",pady=5)
                self.ls_widgets[self.indx].append(self.entry_limit)


                self.bt_limit=CTkButton(self.ls_bts[self.indx],text="Limit",command=self.limit)
                self.bt_limit.grid(row=3,column=0,pady=10)
                self.ls_widgets[self.indx].append(self.bt_limit)

                self.label_abo=CTkLabel(self.ls_bts[self.indx],text="")
                self.label_abo.grid(row=4,column=0,pady=5)
                self.ls_widgets[self.indx].append(self.label_abo)


            
            
                t1=Thread(target=self.qrcodesccaner)
                t1.start()
       except:
           pass
            

    def limit(self):
        try:
            ls=self.q.split("_")
            lim=self.entry_limit.get()
            self.entry_limit.delete(0,END)
            self.q="go"
            data={}
            if os.path.exists(r"F:\goads.json"):
                file=open(r"F:\goads.json")
                data=json.load(file)
                file.close()

                if len(data)!=0:
                    ls_codes=list(data.keys())

                    if ls[0] in ls_codes:

                        for i in range(len(ls_codes)):
                            if ls_codes[i]==ls[0]:
                                if len(data[ls[0]])==3:
                                    data[ls[0]].append(lim)
                                else:
                                    data[ls[0]][-1]=lim
                        with open (r"F:\goads.json","w") as f:
                            json.dump(data,f)

                else:
                    self.label_abo.configure(text="ADD the goad first !!")


                self.entry_code.delete(0,END)
                self.label_abo.configure(text="DONE")


            else:
                self.entry_code.delete(0, END)
                self.label_abo.configure(text="ADD the goad at the first !")
            t1 = Thread(target=self.qrcodesccaner)
            t1.start()
        except:
            pass
        
    def creating_fatora_com(self):
        # self.indx=4
        # # print(self.indx)
        # data={}
        try:
            self.q="go"
            if self.indx!=4:
                for i in self.ls_widgets[self.indx]:
                    i.grid_forget()
                self.indx=4
                # print(self.indx)
                for i in range(11):
                        self.ls_bts[i].grid_forget()

                self.ls_bts[self.indx].grid(row=1,column=0,sticky="nswe")

                self.ls_bts[self.indx].grid_columnconfigure(0,weight=1)

                self.label_title=CTkLabel(self.ls_bts[self.indx],text=" Creating the Fatora ")
                self.label_title.grid(row=0,column=0,sticky="we",pady=5)
                self.ls_widgets[self.indx].append(self.label_title)

                self.entry_code=CTkEntry(self.ls_bts[self.indx],placeholder_text="Scan the goad to get the code",textvariable=self.v_code)
                self.entry_code.grid(row=1,column=0,sticky="we",pady=5)
                self.ls_widgets[self.indx].append(self.entry_code)

                self.entry_amount=CTkEntry(self.ls_bts[self.indx],placeholder_text="Enter the amount")
                self.entry_amount.grid(row=2,column=0,sticky="we",pady=10)
                self.ls_widgets[self.indx].append(self.entry_amount)

                self.name_user=CTkEntry(self.ls_bts[self.indx],placeholder_text="Enter the name of the client")
                self.name_user.grid(row=3,column=0,sticky="we",pady=10)
                self.ls_widgets[self.indx].append(self.name_user)

                self.frame_bts=CTkFrame(self.ls_bts[self.indx],fg_color="transparent")
                self.frame_bts.grid(row=4,column=0,sticky="we",pady=10)
                self.ls_widgets[self.indx].append(self.frame_bts)

                self.frame_bts.grid_rowconfigure(0,weight=1)
                self.frame_bts.grid_columnconfigure((0,2),weight=1)
                # self.frame_bts.grid_columnconfigure(1,weight=1)
                self.bt1=CTkButton(self.frame_bts,text="ADD",command=self.add_goad_fatora)
                self.bt1.grid(row=0,column=0,sticky="e",padx=10)


                self.bt2=CTkButton(self.frame_bts,text="Create the fatora",command=self.create_fatora)
                self.bt2.grid(row=0,column=2,sticky="w",padx=10)

                self.label_done=CTkLabel(self.ls_bts[self.indx],text="")
                self.label_done.grid(row=5,column=0)
                self.ls_widgets[self.indx].append(self.label_done)

                t1=Thread(target=self.qrcodesccaner)
                t1.start()
        except:
            pass
            

            

    def add_goad_fatora(self):
        try:
            self.q="go"
            data_info={}
            code=self.entry_code.get()
            amount=self.entry_amount.get()
            self.name=self.name_user.get()
            indx=0
            if os.path.exists(r"F:\goads.json"):
                with open(r"F:\goads.json","r") as f:
                    data_info = json.load(f)
                ls_info=list(data_info.values())
                ls_codes=list(data_info.keys())
                if code!="" and amount!="" and self.name!="":


                    if code in data_info:
                        indx = ls_codes.index(code)
                        if len(ls_info[indx])==4:
                            if ls_info[indx][0]-int(amount)>=int(ls_info[indx][3]):
                                if len(self.ls_fatora)==0:
                                    self.ls_fatora[self.name]=[[ls_info[indx][2],ls_info[indx][1],amount,(int(amount)*int(ls_info[indx][1]))]]
                                else:
                                    self.ls_fatora[self.name].append([ls_info[indx][2],ls_info[indx][1],amount,(int(amount)*int(ls_info[indx][1]))])
                                data_info[code][0]= ls_info[indx][0]-int(amount)
                                self.label_done.configure(text=" Done ")
                                self.tot_price=self.tot_price+(int(amount)*int(ls_info[indx][1]))
                            else:
                                self.label_done.configure(text="Error : the required amount > allowed amount")

                        else:
                            if ls_info[indx][0]-int(amount)>=0:
                                if len(self.ls_fatora)==0:
                                    self.ls_fatora[self.name]=[[ls_info[indx][2],ls_info[indx][1],amount,(int(amount)*int(ls_info[indx][1]))]]
                                else:
                                    self.ls_fatora[self.name].append([ls_info[indx][2],ls_info[indx][1],amount,(int(amount)*int(ls_info[indx][1]))])
                                data_info[code][0]= ls_info[indx][0]-int(amount)
                                self.label_done.configure(text="Done")
                                self.tot_price=self.tot_price+(int(amount)*int(ls_info[indx][1]))
                            else:
                                self.label_done.configure(text="Error : the required amount > existed amount")
                    else:
                        self.label_done.configure(text=" Add the goad first ")

                else:
                    self.label_done.configure(text="Please , Scan the goads and enter the name of the client !")

                with open(r"F:\goads.json","w") as f:
                    json.dump(data_info,f)
            else:
                self.label_done.configure(text=" Add the goads to your data ")

            t1=Thread(target=self.qrcodesccaner)
            t1.start()
        except:
            pass

    def create_fatora(self):
        try:
            if not(self.create):

                self.create=True
                self.name_user.delete(0,END)
                self.time=time.asctime()
                data={}
                data_info={}
                ind=0
                self.q="1_mohammad" #assumption
                if len(self.ls_fatora)==0:
                    self.label_done.configure("Scan any goad to create !!")
                else:

                    if os.path.exists(r"F:\fatora.json"):
                        with open(r"F:\fatora.json","r") as f:
                            data=json.load(f)
                        if self.name in list(data.keys()):

                            data[self.name].append((list(self.ls_fatora.values()))[0])
                        else:
                            data[self.name]=[(list(self.ls_fatora.values()))[0]]

                        ind=data[self.name].index((list(self.ls_fatora.values()))[0])
                        data[self.name][ind].append(self.tot_price)
                        data[self.name][ind].append(self.time)

                    else:
                        data[self.name]=[(list(self.ls_fatora.values()))[0]]
                        data[self.name][ind].append(self.tot_price)
                        data[self.name][ind].append(self.time)
                    with open (r"F:\fatora.json","w") as f:
                        json.dump(data,f)

                    with open(r"F:\goads.json","r") as f:
                        data_info = json.load(f)
                    ls_info=list(data_info.values())
                    ls_codes=list(data_info.keys())
                    doc = docx.Document()
                    doc.add_heading("Fatora (ELprince Shop)", 0)
                    doc_para = doc.add_paragraph("")
                    doc_para.add_run(f'name :{self.name} \n\ntime:{self.time}').bold=True
                    table = doc.add_table(rows = len(self.ls_fatora[self.name])-2 , cols = 4)
                    table.style='Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Item'
                    hdr_cells[1].text = 'price'
                    hdr_cells[2].text = 'amount'
                    hdr_cells[3].text = 'the price of this amount'

                    print(self.ls_fatora[self.name])
                    for j in range(0,len(self.ls_fatora[self.name])-2):

                        row=table.add_row().cells
                        doc_para.add_run("\n")
                        row[0].text=str(self.ls_fatora[self.name][j][0])
                        row[1].text=str(self.ls_fatora[self.name][j][1])
                        row[2].text=str(self.ls_fatora[self.name][j][2])
                        row[3].text=str(self.ls_fatora[self.name][j][3])
                    doc_para2 = doc.add_paragraph("")
                    doc_para2.add_run(f'\nThe Tot Price Is {self.tot_price}'+"\n").bold=True
                    doc_para2.add_run("\nThank You For Deading With Us , Best Wishs !!").bold=True
                    doc.save(r'F:\fatora.docx')
                    self.ls_fatora={}
                    self.create_earns_json()
                    data={}
                    if os.path.exists(r"F:\j.json"):
                        with open(r"F:\j.json","r") as f:
                            data=json.load(f)
                    if self.name in data:
                        # if data[self.name]>=self.tot_price:
                        data[self.name]=data[self.name]-self.tot_price
                        if data[self.name]>0:
                            self.label_done.configure(text=" D o n e  !!")
                        else:
                            self.label_done.configure(text=f"D o n e !!! , but you must pay a sum of money which is {abs(data[self.name])}")

                        with open(r"F:\j.json","w") as f:
                            json.dump(data,f)
                    else:
                        self.label_done.configure(text=" D o n e  !!")

            self.tot_price=0

            self.create=False
            Thread(target=self.qrcodesccaner).start()
        except:
            pass
        



    def back_goad(self):
        # self.indx=5
        # # print(self.indx)
        # pass
        try:
            self.q="go"
            if self.indx!=5:
                for i in self.ls_widgets[self.indx]:
                    i.grid_forget()
                self.indx=5
                # print(self.indx)
                for i in range(11):
                        self.ls_bts[i].grid_forget()

                self.ls_bts[self.indx].grid(row=1,column=0,sticky="nswe")
                self.ls_bts[self.indx].grid_columnconfigure(0,weight=1)



                self.label_title=CTkLabel(self.ls_bts[self.indx],text="Back any goad")
                self.label_title.grid(row=0,column=0)
                self.ls_widgets[self.indx].append(self.label_title)

                self.entry_name_client=CTkEntry(self.ls_bts[self.indx],placeholder_text="Enter the name of the client ")
                self.entry_name_client.grid(row=1,column=0,sticky='we',pady=10,padx=100)
                self.ls_widgets[self.indx].append(self.entry_name_client)

                self.entry_code=CTkEntry(self.ls_bts[self.indx],placeholder_text="Scan The code here",textvariable=self.v_code)
                self.entry_code.grid(row=2,column=0,sticky="we",pady=10,padx=100)
                self.ls_widgets[self.indx].append(self.entry_code)

                self.entry_amount=CTkEntry(self.ls_bts[self.indx],placeholder_text="Enter the amount here")
                self.entry_amount.grid(row=3,column=0,sticky="we",pady=10,padx=100)
                self.ls_widgets[self.indx].append(self.entry_amount)


                self.bt_back=CTkButton(self.ls_bts[self.indx],text="Go",command=self.back_com)
                self.bt_back.grid(row=4,column=0,sticky="we",pady=20,padx=100)
                self.ls_widgets[self.indx].append(self.bt_back)

                self.label_pro=CTkLabel(self.ls_bts[self.indx],text="")
                self.label_pro.grid(row=5,column=0)
                self.ls_widgets[self.indx].append(self.label_pro)

                t1=Thread(target=self.qrcodesccaner)
                t1.start()
        except:
            pass
    def back_com(self):
        try:
            if not(self.back):

                self.back=True
                self.q="go"
                data={}
                data_goads={}
                self.name=self.entry_name_client.get()
                if os.path.exists(r"F:\goads.json"):
                    with open(r"F:\goads.json","r") as file:
                        data_goads=json.load(file)
                if os.path.exists(r"F:\fatora.json"):
                    with open (r"F:\fatora.json","r") as f:
                        data=json.load(f)
                ls_names=list(data.keys())
                ls_fat_in=list(data.values())
                req_ls=[]
                ls_in=[]
                ind=0
                ls_ele=[]
                amount_fat=0

                if self.v_code.get() in data_goads:
                    if self.entry_name_client.get() in ls_names:
                        ind=ls_names.index(self.entry_name_client.get())
                        req_ls=ls_fat_in[ind]
                        for i in range(len(req_ls)):
                            ls_in=req_ls[i]
                            for j in range(len(ls_in)-2):
                                print(ls_in[j])
                                # print(self.v_name.get())
                                if self.v_name.get() == ls_in[j][0]:
                                    ls_ele = ls_in[j]
                                    print(ls_in[j][2])
                                    amount_fat = amount_fat + int(ls_in[j][2])
                                    print(ls_in[j])

                    # print(amount_fat)
                    # else:
                    #     amount_fat=0

                    # print(amount_fat)

                    # print(amount_fat)

                    if amount_fat==0:

                        self.label_pro.configure(text="you have never bought this goad from us")

                    else:
                        self.b_price = int(self.entry_amount.get()) * ls_ele[1]
                        print(self.b_price)
                        if int(self.entry_amount.get())<=amount_fat:

                            data_goads[self.v_code.get()][0]=data_goads[self.v_code.get()][0]+int(self.entry_amount.get())
                            with open(r"F:\goads.json","w") as f:
                                json.dump(data_goads,f)
                            tot_am=int(self.entry_amount.get())
                            while tot_am!=0:

                                for i in range(len(data[self.name])):
                                    for j in range(len(data[self.name][i])-2):
                                        if data[self.name][i][j][0]==self.v_name.get():
                                            if int(data[self.name][i][j][2])>tot_am:
                                                data[self.name][i][j][2]=str(int(data[self.name][i][j][2])-int(self.entry_amount.get()))
                                                data[self.name][i][j][-1]=data[self.name][i][j][-1]-self.b_price

                                                tot_am=0
                                            else:
                                                tot_am = tot_am - int(data[self.name][i][j][2])

                                                data[self.name][i][j][-1]=0
                                                data[self.name][i][j][2]=str(0)
                                            data[self.name][i][len(data[self.name][i]) - 2] = data[self.name][i][len(data[self.name][i]) - 2] - self.b_price

                            with open (r"F:\fatora.json","w") as f:
                                json.dump(data,f)

                            # t2=Thread(target=self.create_earns_json)
                            # t2.start()
                            data_ac={}
                            self.create_json_not()
                            self.create_earns_json()
                            if os.path.exists(r"F:\j.json"):

                                with open (r"F:\j.json","r") as f:
                                    data_ac=json.load(f)


                            if self.name in data_ac:

                                data_ac[self.name]=data_ac[self.name]+self.b_price

                                with open(r"F:\j.json","w") as f:
                                    json.dump(data_ac,f)
                                self.label_pro.configure(text=" Done , the sum is drawn ! ")
                            else:
                                self.label_pro.configure(text="DONE !!")
                        else:
                            self.label_pro.configure(text="you haven't take all this quantity from us !! ")
                else:
                    self.label_pro.configure(text="Error , this goads isn't from here ! ")





            self.b_price=0
            self.back=False
        
            t1=Thread(target=self.qrcodesccaner)
            t1.start()
        except:
            pass
        
        

    def newuser_com(self):
         try:
             if self.indx!=6:
                for i in self.ls_widgets[self.indx]:
                    i.grid_forget()
                self.indx=6
                # print(self.indx)
                for i in range(11):
                        self.ls_bts[i].grid_forget()

                self.ls_bts[self.indx].grid(row=1,column=0,sticky="nswe")
                self.ls_bts[self.indx].grid_rowconfigure(0,weight=0)
                self.ls_bts[self.indx].grid_rowconfigure((1,2,3,4),weight=1)
                self.ls_bts[self.indx].grid_columnconfigure(0,weight=1)
                self.sign_up_frame=CTkFrame(self.ls_bts[self.indx])
                self.sign_up_frame.grid(row=0,column=0,sticky="we",padx=10)
                self.sign_up_frame.grid_columnconfigure((1,2),weight=0)
                self.sign_up_frame.grid_columnconfigure(0,weight=1)
                self.sign_up_frame.grid_columnconfigure(3,weight=1)
                self.sign_up_frame.grid_rowconfigure(0,weight=1)
                self.ls_widgets[self.indx].append(self.sign_up_frame)
                self.label_signup=CTkLabel(self.sign_up_frame,text="Add a new account",font=("arial",14,"bold"))
                self.label_signup.grid(row=0,column=1,pady=10,padx=10)
                self.ls_widgets[self.indx].append(self.label_signup)

                self.entry_name=CTkEntry(self.ls_bts[self.indx],placeholder_text="Enter the name of the account")
                self.entry_name.grid(row=1,column=0,sticky="we",padx=100,pady=20)
                self.ls_widgets[self.indx].append(self.entry_name)

                self.entry_sum_of_money=CTkEntry(self.ls_bts[self.indx],placeholder_text="Enter the sum of money")
                self.entry_sum_of_money.grid(row=2,column=0,sticky="we",padx=100,pady=20)
                self.ls_widgets[self.indx].append(self.entry_sum_of_money)

                self.bt_add_person=CTkButton(self.ls_bts[self.indx],text=" Add a new person ",font=("arial",14,"bold"),command=self.bt_add_new_person)
                self.bt_add_person.grid(row=3,column=0)
                self.ls_widgets[self.indx].append(self.bt_add_person)

                self.label_done=CTkLabel(self.ls_bts[self.indx],text="",font=("arial",14,"bold"))
                self.label_done.grid(row=4,column=0)
                self.ls_widgets[self.indx].append(self.label_done)
         except:
            pass
          

    def add_money(self):
        try:
            name=self.entry_name_acc
            file=open(r"F:\j.json")
            data = json.load(file)
            ls_names=list(data.keys())
            ls_money=list(data.values())
            name=self.entry_name_acc.get()
            mon=self.entry_money.get()
            if name=="" or mon=="":
                self.label.configure(text="Enter all data")
            else:
                if name in ls_names and mon.isnumeric() :

                    file=open(r"F:\j.json")
                    data = json.load(file)
                    data[name]=float(mon)+float(data[name])
                    with open (r"F:\j.json","w") as f:
                        json.dump(data,f)
                    file.close()
                    self.label.configure(text="Done")

                else:


                    self.label.configure(text="Please , Enter a correct name")
                
        except:
            pass
    
            

        
                     
       
    def add_money_for_account(self):
        try:
            self.new_win=CTkToplevel()
            self.new_win.geometry("400x300")
            self.new_win.title("adding sum of money to the account")
            self.entry_name_acc=CTkEntry(self.new_win,placeholder_text="Enter the name of the account",width=200)
            self.entry_name_acc.pack(pady=10)
            self.entry_money=CTkEntry(self.new_win,placeholder_text="Enter the amount of added money",width=200)
            self.entry_money.pack(pady=10)
            self.bt_go=CTkButton(self.new_win,text="ADD",command=self.add_money)
            self.bt_go.pack(pady=5)
            self.label=CTkLabel(self.new_win,text="")
            self.label.pack()
        except:
            pass


    def bt_add_new_person(self):
        try:
            print(self.indx)

            name=self.entry_name.get()
            sum_mon=self.entry_sum_of_money.get()

            if name=="" or sum_mon=="":

                self.label_done.configure(text=" Enter all data correctly !! ")
                self.label_done.grid(row=4,column=0)

            else:
                ls=name.split(" ")
                print(ls)
                ach = [x.isalpha() for x in ls]
                bol = np.array(ach).all()
                print(bol)
                if bol and sum_mon.isnumeric() :

                    self.write_append_json_file(name,sum_mon)
                    self.entry_name.delete(0,END)
                    self.entry_sum_of_money.delete(0,END)

                else:

                    self.label_done.configure(text=" Enter correct data !! ")
        except:
            pass
            

    def write_append_json_file(self,name,sum_mon):
        #try:
        try:
            data={}
            if os.path.exists(r"F:\j.json"):
                file=open(r"F:\j.json")
                data = json.load(file)
                file.close()
            data[name]=float(sum_mon)
            print(data)
            with open (r"F:\j.json","w") as f:
                json.dump(data,f)
            
            self.label_done.configure(text="Done")
        except:
            pass
        #except:

        #    self.label_done.configure(text="Enter a correct data")



    def search_fatora_com(self):
        try:
            if self.indx!=7:
                for i in self.ls_widgets[self.indx]:
                    i.grid_forget()
                self.indx=7
                # print(self.indx)
                for i in range(11):
                    self.ls_bts[i].grid_forget()
                self.ls_bts[self.indx].grid(row=1,column=0,sticky="nswe")
                self.ls_bts[self.indx].grid_columnconfigure(0,weight=1)
                self.label_title=CTkLabel(self.ls_bts[self.indx],text="Searching about any fatora")
                self.label_title.grid(row=0,column=0,pady=20,sticky="we")
                self.ls_widgets[self.indx].append(self.label_title)

                self.entry_name_client=CTkEntry(self.ls_bts[self.indx],placeholder_text="Enter the name of the client")
                self.entry_name_client.grid(row=1,column=0,sticky="we",pady=10,padx=100)
                self.ls_widgets[self.indx].append(self.entry_name_client)

                self.bt_show=CTkButton(self.ls_bts[self.indx],text="Show the fatora",command=self.bt_showfatora_com)
                self.bt_show.grid(row=2,column=0,sticky="we",pady=5,padx=100)
                self.ls_widgets[self.indx].append(self.bt_show)

                self.label_show=CTkLabel(self.ls_bts[self.indx],text="",font=("arial",13,"bold"))
                self.label_show.grid(row=3,column=0,sticky="nswe",pady=10)
                self.ls_widgets[self.indx].append(self.label_show)
        except:
            pass

    def bt_showfatora_com(self):
        try:
            data={}
            name="0"
            if name!=self.entry_name_client.get():

                self.label_show.configure(text="")
                name=self.entry_name_client.get()
                if os.path.exists(r"F:\fatora.json"):
                    with open(r"F:\fatora.json","r") as f:
                        data=json.load(f)
                    ls_names=list(data.keys())
                    ls_info=list(data.values())
                    ind=0
                    req_ls=[]
                    ls_infat=[]


                    if name in ls_names:

                        ind=ls_names.index(name)

                        req_ls=ls_info[ind]
                        for j in range(len(req_ls)):
                            ls_infat=req_ls[j]
                            self.label_show.configure(text=self.label_show.cget("text")+"\n---------------------------------\n"+f"name is {name} \n\nTime of this fatora is {ls_infat[-1]} \n\n the total price is {ls_infat[len(ls_infat)-2]}")

                            for i in range (len(ls_infat)-2):

                                self.label_show.configure(text=self.label_show.cget("text")+"\n\n"+f"an amount of {ls_infat[i][0]}  =  {ls_infat[i][2]}  >>>  the price is {ls_infat[i][-1]} $")


                    else:
                        self.label_show.configure(text=" there is no client with this name <fatora is deleted or not created > !! ")
                else:
                    self.label_show.configure(text=" There is no one have bought any thing from us !! ")
        except:
            pass


    def srch_persons_com(self):
        try:
            if self.indx!=8:
                for i in self.ls_widgets[self.indx]:
                    i.grid_forget()
                self.indx=8
                # print(self.indx)
                for i in range(11):
                    self.ls_bts[i].grid_forget()
                self.ls_bts[self.indx].grid(row=1,column=0,sticky="nswe")
                self.ls_bts[self.indx].grid_rowconfigure((0,1,2,3),weight=0)
                self.ls_bts[self.indx].grid_columnconfigure(0,weight=1)

                self.label_srch=CTkLabel(self.ls_bts[self.indx],text=" Searching about any person on the system ",font=("arial",14,"bold"))
                self.label_srch.grid(row=0,column=0)
                self.ls_widgets[self.indx].append(self.label_srch)

                self.entry_name_srch=CTkEntry(self.ls_bts[self.indx],placeholder_text=" Enter the name here ")
                self.entry_name_srch.grid(row=1,column=0,sticky="we",padx=100,pady=20)
                self.ls_widgets[self.indx].append(self.entry_name_srch)


                self.bt_show=CTkButton(self.ls_bts[self.indx],text="Show",command=self.bt_show_com)
                self.bt_show.grid(row=2,column=0,pady=5)
                self.ls_widgets[self.indx].append(self.bt_show)

                self.add=CTkButton(self.ls_bts[self.indx],text="Add Money to the person's account account",command=self.add_money_for_account)
                self.add.grid(row=3,column=0,pady=10)
                self.ls_widgets[self.indx].append(self.add)

                self.label_show=CTkLabel(self.ls_bts[self.indx],text="",font=("arial",13,"bold"))
                self.label_show.grid(row=4,column=0,sticky="we",pady=5)
                self.ls_widgets[self.indx].append(self.label_show)
        except:
            pass
          


    
    def bt_show_com(self):
        try:
            if self.entry_name_srch.get()=="":

                self.label_show.configure(text=" Enter the name ")

            else:


                    name=self.entry_name_srch.get()
                    self.entry_name_srch.delete(0,END)
                    file=open(r"F:\j.json")
                    data = json.load(file)
                    ls_names=list(data.keys())
                    ls_money=list(data.values())

                    if name in ls_names:
                        for i in range(len(ls_names)):

                            if name==ls_names[i]:

                                self.label_show.configure(text=f"the name is {name} \n the sum of money here is {ls_money[i]}")

                    else:

                        self.label_show.configure(text=" the name isn't on the system ")
        except:
            pass
   
                

    def gen_qr_code(self):
        try:
            # self.indx=9
            # pass
            if self.indx!=9:
                for i in self.ls_widgets[self.indx]:
                    i.grid_forget()
                self.indx=9
                # print(self.indx)
                for i in range(11):
                    self.ls_bts[i].grid_forget()
                self.ls_bts[self.indx].grid(row=1,column=0,sticky="nswe")
                self.ls_bts[self.indx].grid_columnconfigure(0,weight=1)

                self.label_t=CTkLabel(self.ls_bts[self.indx],text="Enter the info to create the QR-Code")
                self.label_t.grid(row=0,column=0,pady=10)
                self.ls_widgets[self.indx].append(self.label_t)

                self.entry_name=CTkEntry(self.ls_bts[self.indx],placeholder_text="Enter the name of the goad")
                self.entry_name.grid(row=1,column=0,sticky="we",pady=5)
                self.ls_widgets[self.indx].append(self.entry_name)

                self.entry_loc_save=CTkEntry(self.ls_bts[self.indx],placeholder_text="Enter the location of the saving")
                self.entry_loc_save.grid(row=2,column=0,sticky="we",pady=5)
                self.ls_widgets[self.indx].append(self.entry_loc_save)

                self.bt_gen=CTkButton(self.ls_bts[self.indx],text=" Generate",command=self.press_gen)
                self.bt_gen.grid(row=3,column=0,sticky="we",pady=10)
                self.ls_widgets[self.indx].append(self.bt_gen)

                self.label_done=CTkLabel(self.ls_bts[self.indx],text="")
                self.label_done.grid(row=4,column=0)
                self.ls_widgets[self.indx].append(self.label_done)
        except:
            pass

            
    def press_gen(self):
        try:
            name=self.entry_name.get()
            decod_inf="1"
            self.entry_name.delete(0,END)
            location=self.entry_loc_save.get()
            self.entry_loc_save.delete(0,END)

            data={}

            if not(os.path.exists(location)):

                self.entry_loc_save.configure(placeholder_text="Enter a correct path of saving ")

            else:
                if os.path.exists(r"F:\goads.json"):

                    with open(R"F:\goads.json","r") as f:
                        data=json.load(f)


                decod_inf=str(len(data)+1)+f"_{name}"

                image=np.array(qrcode.make(decod_inf))
                image = Image.fromarray(image) #PIL Image
                img = CTkImage(dark_image=image,size=(250,250))
                # Call draw Method to add 2D graphics in an image
                I1 = ImageDraw.Draw(image)
                myfont=ImageFont.truetype('arial.ttf', 30)
                # Add Text to an image
                I1.text((100,250 ), name,font=myfont, fill=0)

                image.save(location+"\\"+decod_inf+".png")
                self.label_done.configure(text="Done !! ")
        except:
            pass
        
    
        
    def aboutus_com(self):
        try:
            if self.indx!=10:
                for i in self.ls_widgets[self.indx]:
                    i.grid_forget()
                self.indx=10
                # print(self.indx)
                for i in range(11):
                    self.ls_bts[i].grid_forget()
                self.ls_bts[self.indx].grid(row=1,column=0,sticky="nswe")

                self.label_img_com=CTkLabel(self.ls_bts[self.indx],text="",image=img4)
                self.label_img_com.grid(row=0,column=0,pady=20,padx=300,sticky="nswe")
                self.ls_widgets[self.indx].append(self.label_img_com)
                self.label_name_com=CTkLabel(self.ls_bts[self.indx],text=" E l p r i n c e ",font=("arial",15,"bold"))
                self.label_name_com.grid(row=1,column=0,pady=10,padx=300,sticky="nswe")
                self.ls_widgets[self.indx].append(self.label_name_com)
        except:
            pass
###############################################################


        



s=shop()
s.mainloop()




